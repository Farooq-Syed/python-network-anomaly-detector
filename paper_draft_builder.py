"""Build an anonymized workshop-paper draft from the frozen PNAD results."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "PNAD_anonymous_workshop_paper.docx"


def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_mar = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                cell._tc.get_or_add_tcPr().append(tc_mar)
            for side in ("top", "start", "bottom", "end"):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "75" if side in ("top", "bottom") else "100")
                node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    """Mark the first row as a repeating header for accessibility and pagination."""
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1); run._r.append(instr); run._r.append(fld_char2)
    set_font(run, 9)


def add_text(doc, text, style="Normal", align=None, before=None, after=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    set_font(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r, 12 if level == 1 else 11, bold=True)
    return p


def add_table(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value); set_font(r, font_size, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, (cell, value) in enumerate(zip(cells, row)):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value)); set_font(r, font_size)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text); set_font(r, 8.5, italic=True)


def add_table_reading(doc, text):
    """Give a non-specialist one concrete row-level interpretation."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.03
    r = p.add_run("How to read this table. ")
    set_font(r, 8.8, bold=True, color=(36, 74, 115))
    r = p.add_run(text)
    set_font(r, 8.8)
    return p


def add_figure(doc, image_path, width, alt_text):
    """Insert a centered figure and attach meaningful alternative text."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(image_path, width=width)
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", "Benchmark collection and offline evaluation architecture")
    return p


def add_reference(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    p.paragraph_format.space_after = Pt(0.5)
    p.paragraph_format.line_spacing = 0.95
    r = p.add_run(text); set_font(r, 7.8)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.68)
section.left_margin = Inches(0.82)
section.right_margin = Inches(0.82)
section.header_distance = Inches(0.32)
section.footer_distance = Inches(0.32)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.08

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
header.paragraph_format.space_after = Pt(0)
r = header.add_run("Anonymous Submission")
set_font(r, 9, italic=True)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.paragraph_format.space_before = Pt(0)
r = footer.add_run("Page ")
set_font(r, 9)
add_page_field(footer)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(4)
title.paragraph_format.space_after = Pt(5)
r = title.add_run("When Active Learning Does Not Beat Random Labeling for Network Intrusion Detection")
set_font(r, 15, bold=True)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(10)
r = subtitle.add_run("Anonymous Workshop Submission | Reviewer-Feedback Revision | August 2026")
set_font(r, 10, italic=True)

audience = doc.add_paragraph()
audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
audience.paragraph_format.space_before = Pt(0)
audience.paragraph_format.space_after = Pt(8)
r = audience.add_run("Audience and scope: a specialized cybersecurity measurement paper, written to remain accessible to general cybersecurity reviewers; ML is used as an evaluation method, not presented as a new architecture.")
set_font(r, 9, bold=True, color=(36, 74, 115))

add_heading(doc, "Abstract", 1)
add_text(doc, "Security analysts cannot label every network-flow record, so active learning attempts to reduce their workload by asking for labels only on selected records. This paper tests whether that selection is actually better than choosing records at random. We compare random labeling with posterior uncertainty, diversity sampling, and query-by-committee on prepared subsets of three public intrusion-detection benchmarks: UNSW-NB15, CIC-IDS2017, and CSE-CIC-IDS2018. Performance is summarized primarily by F1, the harmonic mean of precision and recall, on a 0-to-1 scale where higher is better. On CIC-IDS2017, every non-random strategy is significantly worse than random labeling at a 100-label budget after correction for multiple comparisons (F1 0.797-0.846 versus 0.889) with the original logistic learner. On the more separable UNSW-NB15 subset, no alternative remains significantly better than random. Stricter tests also show that supervised F1 near 0.97 under balanced random cross-validation falls sharply when entire days or datasets are held out; on a 5% reweighted CIC subset, recall is only 0.579-0.665 when thresholds maximize validation recall subject to a 1% false-positive budget. This is an offline cybersecurity benchmark study, not a new classifier or a deployment trial. Its narrow result concerns the evaluated policies, learners, subsets, and budgets; it is not a verdict on active learning for NIDS generally.")

add_heading(doc, "1. Introduction", 1)
add_text(doc, "Network intrusion detection must distinguish harmful activity from a large and changing background of legitimate traffic. Labels can make that task much easier, but labeling flows is expensive and often incomplete. Active learning offers an appealing response: ask a human to label only the flows that appear most informative. In security practice, however, queried examples may be ambiguous because of noise, overlap, or collection artifacts rather than because they reduce uncertainty about the deployed decision boundary.")
add_text(doc, "This paper studies a focused empirical question: when a modest label budget is available for flow-based intrusion detection, does active selection beat random labeling? We evaluate a simple, transparent logistic-regression base learner and four selection policies: random sampling, posterior uncertainty, diversity/representativeness sampling, and query-by-committee (QBC). For binary logistic regression, margin and entropy sampling induce the same ranking as posterior uncertainty, so we intentionally report them as one baseline rather than inflate the comparison set.")
add_text(doc, "The central result is negative. On the harder CIC-IDS2017 subset, all three active strategies are significantly worse than random labeling at the fixed budget. On easy UNSW-NB15 data, the strategies are statistically indistinguishable at the performance ceiling. We then ask whether the attractive random-CV results survive stronger tests. They do not consistently survive attack-family, day, cross-dataset, and operating-point shifts. These findings make a limited claim: active querying is not a free improvement for the prepared datasets and protocols studied here.")
add_text(doc, "The intended audience is cybersecurity, intrusion-detection, and security-ML researchers who understand basic classification but may not specialize in active learning. Machine learning is therefore used as an evaluation instrument for a network-defense question, not presented as a new model architecture. The next subsection defines the principal records, metrics, and workflow before presenting the technical protocol.")

add_heading(doc, "1.1 Reader guide: records, metrics, and labeling", 2)
add_text(doc, "A network flow is a summary of communication between endpoints, such as duration, bytes, packets, ports, and timing statistics; it is not the packet payload itself. A positive prediction means that the model flags a flow as malicious. Precision is the fraction of flagged flows that are attacks, recall is the fraction of attacks that are detected, and F1 is their harmonic mean. F1 ranges from 0 to 1 and is high only when both precision and recall are high. The false-positive rate (FPR) is the fraction of benign flows incorrectly flagged. In operational terms, false positives consume analyst attention, which is why the paper reports both F1 and recall at a validation-selected FPR target.")
add_text(doc, "A label budget is the maximum number of training records for which the experiment reveals the benign/attack label, standing in for limited analyst annotation. Random sampling selects those records without using the current model. Posterior uncertainty selects records whose predicted attack probability is closest to 0.5. Diversity sampling seeks coverage of different regions of the feature space. Query-by-committee selects records on which several models disagree. Every policy receives the same budget, initial labeled seed, learner, split, and random seed so that the acquisition rule is the intended difference.")

add_heading(doc, "1.2 Why this comparison is needed", 2)
add_text(doc, "Classical active learning aims to improve a learner with fewer labels by choosing examples that are expected to be especially informative. Uncertainty sampling is the most direct version of this idea: label the point near the model's current boundary. Diversity policies add a coverage intuition by seeking points that represent distinct regions of the pool. Committee methods substitute disagreement among plausible models for the confidence of a single model. These ideas are well established, and each can be sensible when the unlabeled pool resembles the eventual test distribution and labels are consistent enough for boundary refinement to be useful.")
add_text(doc, "Network-flow benchmarks make those assumptions difficult to inspect. Attack labels can group heterogeneous behavior; benign traffic can shift with host, application, collection day, or capture environment; and a detector may face a prevalence far below the balanced research subset. A method that selects boundary cases can therefore spend its scarce labels on examples that are ambiguous in the chosen feature space rather than informative for the target task. Conversely, random labeling may cover common regions of the pool surprisingly well when the model is simple and the benchmark already separates benign and malicious traffic.")
add_text(doc, "Prior intrusion-detection studies commonly report high random-split accuracy or F1, and active-learning papers often compare several query curves on a single benchmark partition. Those studies are valuable baselines, but they do not by themselves answer whether a selection rule gives a statistically reliable benefit at a fixed small budget or whether the resulting detector withstands a collection shift. The present work is structured around those two questions. It deliberately uses transparent models and standard policies so that the negative result cannot be attributed to an opaque architecture or an unmatched implementation.")
add_text(doc, "The objective is not to rule out active learning. A different labeler, feature representation, pool, or cost function can alter the comparison. Instead, the contribution is a falsifiable reference point: under matched budgets and repeated seeds, the commonly used policies tested here do not produce a dependable gain over random selection. This distinction matters when a claimed improvement is intended to save analyst effort rather than merely improve one retrospective learning curve.")

add_heading(doc, "Contributions", 2)
for item in [
    "A controlled comparison of four genuinely distinct label-acquisition policies under identical budgets, splits, features, and base learner.",
    "Repeated-seed paired tests showing that each non-random policy is significantly worse than random labeling on CIC-IDS2017, while no policy is significantly better on UNSW-NB15.",
    "A stricter-evaluation package that separates in-distribution random-CV performance from unseen-family, whole-day, cross-dataset, and reweighted operating-point evidence.",
    "A reproducible artifact with fixed seeds, metadata-preserving subsets, test coverage, and exact commands for each reported result."
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.04
    r = p.add_run(item); set_font(r)

add_heading(doc, "2. Experimental Design", 1)
add_heading(doc, "2.1 Data and preprocessing", 2)
add_text(doc, "The study uses three public flow-based benchmark families. Each prepared benchmark subset is balanced for the primary comparisons and retains only numeric traffic features; metadata-retaining variants preserve attack family or day where needed for strict splits. The work does not treat these curated subsets as operational traces. In particular, the 5% experiment below is a reweighted subset, not a naturally observed deployment prevalence.")
add_text(doc, "No new traffic was captured for this study, and the author did not construct a local router, host, sensor, or attack testbed. The benchmark providers created the underlying collection environments; this work begins with their released flow tables and evaluates them offline. Consequently, the experimental architecture to reproduce here is the data and information flow below, not a physical equipment diagram. This distinction also prevents the public benchmarks from being mistaken for evidence about a specific enterprise network.")
add_table(doc,
    ["Dataset", "Primary subset", "Strict metadata", "Purpose"],
    [
        ["UNSW-NB15", "12,000 balanced rows; 21 numeric features", "attack family", "easy/separable reference; unseen-family tests"],
        ["CIC-IDS2017", "12,000 balanced rows; 78 numeric features", "day and attack name", "hard active-learning and temporal-shift tests"],
        ["CSE-CIC-IDS2018", "12,000 balanced rows; CICFlowMeter schema", "prepared subset", "additional benchmark and cross-dataset transfer"],
    ], [2050, 2700, 1700, 2910], 8.1)
add_caption(doc, "Table 1. Prepared data used in the study. All headline limitations are retained in the evaluation and discussion.")
add_table_reading(doc, "For example, the CIC-IDS2017 row means that the study uses 12,000 prepared flow records described by 78 numeric traffic features. Day and attack-name fields are retained only to construct stricter test splits; this dataset supplies the main active-learning and day-shift evidence.")

add_heading(doc, "2.2 Benchmark collection context and evaluation architecture", 2)
add_text(doc, "The three datasets originate from controlled provider testbeds, but their collection environments differ. UNSW-NB15 was generated in the UNSW Canberra Cyber Range using IXIA PerfectStorm for a mixture of normal and synthetic attack behavior; tcpdump captured packets and Argus/Bro-based processing contributed the released features [1]. CIC-IDS2017 used separate victim and attack networks, 12 heterogeneous victim machines, mirrored traffic capture, and CICFlowMeter flow extraction [2]. CSE-CIC-IDS2018 scaled the profile-based approach to 50 attacking machines and a victim organization containing five departments, 420 machines, and 30 servers; CICFlowMeter-V3 produced its flow features [3]. These are benchmark-provider architectures, not equipment built or operated by the author.")
add_text(doc, "This study begins at the released-flow-table boundary shown in Figure 1. It neither replays the attacks nor reconstructs the physical networks. Metadata required for a family or day holdout is retained only to construct the split and is not supplied as a predictive feature. Training and validation labels may fit models or select thresholds, whereas outer-test labels remain hidden until final scoring. During active-learning simulation, only labels selected within the fixed budget are revealed to the acquisition policy.")
add_figure(doc, "assets/evaluation_architecture.png", Inches(6.55), "Three provider testbeds feed released flow tables into the offline study. The study prepares numeric features, partitions train, validation, and untouched test data, acquires labels within a fixed budget, fits and calibrates the model, and evaluates once on the test fold. Packet payloads, equipment state, host telemetry, and analyst context are outside the study boundary.")
add_caption(doc, "Figure 1. Provider collection contexts and the author-controlled offline evaluation boundary. Counts and capture details follow the official dataset descriptions [1-3].")

add_heading(doc, "2.3 Models and acquisition policies", 2)
add_text(doc, "The unsupervised reference is a vote of z-score, Isolation Forest, Local Outlier Factor, and one-class SVM detectors. The label-budget and active-learning experiments use StandardScaler followed by logistic regression, with class balancing where stated. Each acquisition run begins from a small seed set containing both classes, then requests labels in batches until reaching the specified budget. The same train/test folds and seeds are used for every strategy.")
add_text(doc, "Posterior uncertainty selects points closest to 0.5 under the current classifier. Diversity sampling chooses representative points near clusters in the unlabeled pool. QBC trains a small committee on random 80% subsets of the labeled pool, without replacement, and ranks points by normalized entropy of the members' hard votes. Thus a confident 50/50 committee split is maximally uncertain, as required by the QBC interpretation.")

add_heading(doc, "2.4 Evaluation protocol", 2)
add_text(doc, "Primary scores use stratified five-fold cross-validation. The active-learning comparison reports per-seed F1 at a 100-label budget across repeated seeds, then uses paired two-sided Wilcoxon signed-rank tests with Bonferroni correction. This design answers whether observed differences are larger than run-to-run variation, not simply whether one curve appears higher in one split.")
add_text(doc, "We also include stricter tests. Family holdout removes an entire UNSW attack family from supervised training while retaining a benign test split. Day holdout places every CIC flow from a held-out day, benign and attack, in the test set. Cross-dataset transfer trains on the shared CICFlowMeter feature intersection and evaluates on CSE-CIC-IDS2018. Finally, the 5% operating-point frame is constructed deterministically before cross-validation by retaining every benign row and sampling attacks. Within each outer training fold, a stratified 5-fold inner partition reserves 20% for validation. The model is fitted on the remaining 80%; the threshold maximizes validation recall subject to FPR no greater than 1% and is applied once to the untouched outer test fold.")

add_heading(doc, "2.5 Protocol choices that prevent optimistic comparisons", 2)
add_text(doc, "Several implementation choices matter for interpretation. For binary logistic regression, least-confidence, smallest-margin, and entropy rules are monotone transformations of the same posterior uncertainty ranking. Reporting all of them as separate successful methods would count one behavior multiple times, so this study keeps one posterior-uncertainty policy and adds QBC as a substantively different policy. The artifact includes a test that verifies the ranking equivalence for the binary posterior case.")
add_text(doc, "Likewise, QBC disagreement is computed from hard committee votes rather than from the average distance of member probabilities from 0.5. The latter can incorrectly regard a committee split between confident benign and confident attack predictions as certain. Vote entropy instead gives that split the highest disagreement score and gives unanimous votes a score near zero. Committee members train on independently drawn random 80% subsets of the current labeled pool without replacement; this is a lightweight source of model variation, not bootstrap resampling.")
add_text(doc, "The strict day protocol guards against a more subtle form of evaluation optimism. If only attacks from a held-out day are removed while benign examples from the same day remain in training, the test is not a complete day shift. Here the full held-out day is test-only. Similarly, recall at a false-positive target is not computed by choosing the best threshold on the test fold. Each outer fold reserves an inner validation split for threshold choice, and the resulting threshold is applied once to the untouched outer test data.")
add_text(doc, "These controls do not make public benchmarks identical to live networks. They do make the claims easier to audit. A reader can disagree with the choice of data, learner, or metric without needing to infer whether the reported result was helped by a duplicated query strategy, shared temporal rows, or a threshold tuned after observing the evaluation labels.")

add_heading(doc, "3. Results", 1)
add_heading(doc, "3.1 Labels matter, but the headline can mislead", 2)
add_text(doc, "Table 2 first establishes the familiar result that labels help much more than changing among several standard anomaly detectors. On all three prepared benchmarks, a supervised baseline greatly exceeds the unsupervised vote. This result should not be read as deployment readiness: it is the starting point for the stricter evaluations that follow.")
add_table(doc,
    ["Dataset", "Unsupervised ensemble F1", "Supervised F1", "Difference"],
    [["UNSW-NB15", "0.270", "0.996", "+0.726"], ["CIC-IDS2017", "0.264", "0.966", "+0.702"], ["CSE-CIC-IDS2018", "0.183", "0.924", "+0.741"]],
    [2200, 2350, 2350, 2460], 8.3)
add_caption(doc, "Table 2. Balanced random-CV comparisons. F1 ranges from 0 to 1; these are in-distribution reference results, not deployment estimates.")
add_table_reading(doc, "In the UNSW-NB15 row, the detectors that received no attack labels achieved F1 0.270, whereas the classifier trained with labels achieved 0.996. In plain terms, labels produced far more balanced attack detection on this prepared random split. The +0.726 difference is a benchmark gap, not a promised deployment improvement.")

add_heading(doc, "3.2 Random labeling wins on the hard benchmark", 2)
add_text(doc, "At 100 labels on CIC-IDS2017, random sampling reaches F1 0.889. Posterior uncertainty, diversity, and QBC respectively reach 0.808, 0.797, and 0.846. Each deficit is significant after correction. On UNSW-NB15, all approaches are near the 0.99 ceiling and none survives correction. QBC has the largest UNSW point estimate, but its +0.005 difference is not significant. The result is therefore not that active learning never helps; it is that none of the tested strategies reliably improves upon random labeling under these conditions.")
add_text(doc, "The UNSW ceiling is itself informative rather than uninformative. When a dataset is this separable at the chosen feature and learner, there is no headroom for an acquisition rule to create: the gain comes from the labels themselves, not from which records are picked. A policy comparison therefore cannot discriminate strategies on UNSW precisely because the labeling problem is easy there. The discriminating evidence comes from the harder CIC subset, where the models are not already near-certain and the choice of which records to label measurably changes the outcome.")
add_table(doc,
    ["Dataset", "Strategy", "F1 (95% CI)", "Difference vs random", "Corrected result"],
    [
        ["CIC", "Random", "0.889 [0.882, 0.895]", "-", "baseline"],
        ["CIC", "Uncertainty", "0.808 [0.783, 0.833]", "-0.081", "worse, p=0.023"],
        ["CIC", "Diversity", "0.797 [0.768, 0.827]", "-0.091", "worse, p=0.023"],
        ["CIC", "QBC", "0.846 [0.817, 0.874]", "-0.043", "worse, p=0.047"],
        ["UNSW", "Random", "0.990 [0.987, 0.993]", "-", "baseline"],
        ["UNSW", "QBC", "0.995 [0.995, 0.996]", "+0.005", "not significant"],
    ], [1100, 1520, 2500, 1900, 2340], 7.8)
add_caption(doc, "Table 3. Repeated-seed active-learning statistics at 100 labels. Values are rounded from the frozen results.")
add_table_reading(doc, "The CIC uncertainty row says that selecting the flows closest to the model's decision boundary produced F1 0.808, with a repeated-run 95% interval of 0.783-0.833. That is 0.081 below random labeling; the corrected p=0.023 indicates that the disadvantage remained statistically detectable after accounting for the strategy comparisons.")

add_heading(doc, "3.2.1 Interpreting the negative comparison", 2)
add_text(doc, "The CIC result is stronger than a comparison of two mean scores. Every active policy uses the same label budget, model family, test folds, and repeated seeds as random sampling. The paired analysis therefore tests the acquisition decision itself. The p-values are Bonferroni-corrected, so they are more conservative (less likely to reject the null) than uncorrected tests; the corrected values (0.023, 0.023, 0.047) are modest rather than overwhelming, but all three point in the same direction: selecting uncertain, representative, or committee-disputed flows did not improve the classifier and yielded lower F1 than simply labeling randomly selected flows.")
add_text(doc, "This outcome cautions against a common operational shortcut: treating an acquisition heuristic as a guaranteed multiplier on analyst time. The tested policies may concentrate on boundary cases whose labels are inherently variable, on portions of the data distribution that are poorly represented by the fixed feature set, or on examples that improve probability calibration without improving the F1 operating point. The experiments were not designed to distinguish these mechanisms, so we do not attribute the gap to any one cause. Their shared practical implication is enough: the query policy must be evaluated against random labeling under the actual budget and target metric.")

add_heading(doc, "3.3 Random-CV success does not transfer consistently", 2)
add_text(doc, "The strict tests put the high supervised random-CV figures in context. In UNSW, holding out one attack family still produces high F1 for most families, although estimates for Backdoor, Worms, and Shellcode are unstable because they contain few attacks. In CIC, entire held-out days behave very differently: Friday and Wednesday retain some separation, while Tuesday and Thursday collapse at the default operating point. Training on CIC shared features and testing on CSE-CIC-IDS2018 reduces F1 to 0.006, despite a 0.516 cross-validation reference on the training dataset's shared features. These are distribution-shift measurements, not evidence of information leakage.")
add_table(doc,
    ["Evaluation", "Test setting", "Supervised F1", "AUC / context"],
    [
        ["UNSW family holdout", "10 held-out families", "0.667-0.998", "small families widen uncertainty"],
        ["CIC day holdout", "Friday", "0.620", "AUC 0.897"],
        ["CIC day holdout", "Wednesday", "0.750", "AUC 0.706"],
        ["CIC day holdout", "Thursday / Tuesday", "0.007 / 0.000", "AUC 0.674 / 0.688"],
        ["Cross-dataset", "CIC -> CSE-CIC", "0.006", "AUC 0.574; 27 shared features"],
    ], [1760, 2200, 1900, 3500], 8.0)
add_caption(doc, "Table 4. Generalization tests reveal substantial distribution sensitivity beyond random cross-validation.")
add_table_reading(doc, "The Friday row means that every Friday flow was excluded from training and used only for testing. F1 0.620 shows that the fixed attack decisions were imperfect, while AUC 0.897 shows that the scores still ranked many attacks above benign flows. The two metrics answer different operational questions.")

add_heading(doc, "3.3.1 What the strict scores do and do not show", 2)
add_text(doc, "The strict results should not be collapsed into a single assertion that the detector has no value. Friday and Wednesday retain nontrivial AUC, and several UNSW family holdouts remain strong. Instead, the evidence shows heterogeneity: a random cross-validation score averages over mixtures that can be much easier than a chronological or cross-collection test. A near-zero F1 at one fixed threshold may coexist with ranking information, but it is still a serious warning for an alerting system whose operating threshold was chosen elsewhere.")
add_text(doc, "The cross-dataset experiment is especially conservative because only a 27-feature shared schema is available for CIC-IDS2017 and CSE-CIC-IDS2018. Its F1 of 0.006 and AUC of 0.574 do not isolate the cause of failure: shifted traffic, preprocessing differences, label construction, and the restricted feature intersection can all contribute. The result nevertheless bounds a tempting interpretation of the training benchmark: a model that scores well under a within-dataset random split should not be presented as cross-environment capable without a direct transfer test.")
add_text(doc, "The family-holdout analysis provides a complementary caution. It is closer to a novelty test because an entire named attack family is removed from training, but its small groups reduce statistical precision. We therefore report it as supporting context rather than as a claim that every unseen attack will be detected. The Backdoor, Worms, and Shellcode families contain only a handful of attack rows, so their held-out F1 estimates carry wide uncertainty and would not survive a bootstrap confidence interval at conventional width; we keep them in the table for transparency but do not base any conclusion on them. Across all strict evaluations, the stable conclusion is modest: the apparent gap between supervised and unsupervised performance is conditional on the partition and on the target operating point.")

add_heading(doc, "3.4 A fixed false-positive target changes the conclusion", 2)
add_text(doc, "A classifier deployed into a security workflow is constrained by alert volume. The 5% attack prevalence used here is a deliberate conservative assumption for an enterprise-telemetry-style operating point: realistic traffic is overwhelmingly benign (often 99%+), so 5% is if anything generous to the detector and represents an upper bound on the attack share an alerting system would typically face; it is not chosen from the benchmark's own class mix. On a 5% reweighted CIC subset, balanced class weights maximize recall but produce low precision; unweighted logistic regression trades recall for a higher F1. With thresholds selected in inner validation folds, recall at a 1% validation-FPR budget is 0.579 for the weighted model and 0.665 for the unweighted model. The comparable UNSW values are 0.972 and 0.975. Mean achieved outer-test FPR is 0.0057 on CIC and 0.0075-0.0078 on UNSW. Consequently, the balanced CIC F1 of 0.97 is not an adequate summary of the operational tradeoff.")
add_table(doc,
    ["Dataset", "Model", "Precision", "Recall", "F1", "Recall @ 1% FPR"],
    [
        ["CIC 5%", "balanced weights", "0.393", "0.946", "0.556", "0.579"],
        ["CIC 5%", "unweighted", "0.926", "0.608", "0.733", "0.665"],
        ["UNSW 5%", "balanced weights", "0.868", "0.997", "0.928", "0.972"],
        ["UNSW 5%", "unweighted", "0.873", "0.975", "0.921", "0.975"],
    ], [1300, 1800, 1250, 1250, 1100, 2660], 8.0)
add_caption(doc, "Table 5. Operating-point evaluation on reweighted prepared subsets. Thresholds are chosen without using outer-test labels.")
add_table_reading(doc, "For the CIC balanced-weights row, precision 0.393 means roughly 39% of flagged flows were attacks, while recall 0.946 means roughly 95% of attacks were found at the default decision rule. After selecting the highest-recall validation threshold that stayed within a 1% FPR budget, recall on the untouched test folds fell to 0.579, exposing the alert-burden tradeoff.")

add_heading(doc, "3.5 Learner-family sensitivity", 2)
add_text(doc, "The headline active-learning comparison uses a linear logistic-regression learner. Active learning is often argued to be most useful for flexible, higher-capacity models where uncertainty is meaningful, so a critic could suspect the negative result is an artifact of the linear base learner. To test this, we repeated the matched random-vs-uncertainty comparison at the 100-label budget with a nonlinear histogram gradient-boosting learner (sigmoid-calibrated, with calibration fit only on the currently labeled training pool). At budget 100 on CIC-IDS2017, the direction is unchanged and stronger: random F1 0.917 versus uncertainty F1 0.658, compared with 0.887 versus 0.813 for the linear learner. The uncertainty policy again underperforms random labeling, so the conclusion is not specific to the linear base learner under the tested conditions. This is a bounded sensitivity check (two seeds, one budget) rather than a full learner-family survey; it is reported to bound the learner-specificity concern, not to generalize across all model classes.")

add_heading(doc, "4. Discussion", 1)
add_text(doc, "For a cybersecurity reader, the result is best understood as an evaluation warning rather than an algorithmic contest. The tested ML methods operate on flow summaries after collection; they do not observe payload content, host telemetry, or analyst context. The comparison therefore asks whether a particular way of allocating a scarce labeling budget improves this constrained detector, not whether active learning is universally unsuitable for security operations.")
add_text(doc, "The active-learning failure is plausible for two nonexclusive reasons. First, a benchmark can be easy enough that random labels rapidly anchor an accurate decision boundary; little room remains for query selection to help. Second, difficult or overlapping flows can be disproportionately selected by uncertainty-based policies. The calibration analyses accompanying this artifact are consistent with the latter explanation on CIC, but the present study does not establish a causal decomposition of noise, overlap, and calibration.")
add_text(doc, "The evaluation also changes how the supervised-versus-unsupervised contrast should be interpreted. Labels provide a large gain within the prepared datasets, but that fact alone does not show that a model will transfer across collection days or benchmarks. The paper's value is therefore methodological: report the random-CV baseline, but pair it with conditions that more closely expose how much of that baseline depends on the data distribution and the thresholding protocol.")
add_heading(doc, "4.1 Implications for evaluation and deployment", 2)
add_text(doc, "For researchers, the first implication is comparative rather than prescriptive. A new query strategy should be tested against random labeling with paired repetitions, fixed budgets, and a stated multiplicity correction when several strategies are compared. Reporting only the best seed or an area under a learning curve can obscure whether a practical advantage remains at the label budget an analyst can afford. A negative benchmark result is also useful: it identifies a setting in which a proposed efficiency gain did not materialize and supplies a reproducible baseline for methods that claim to address it.")
add_text(doc, "For practitioners, the second implication is to separate ranking performance from the alerting decision. The day-holdout results retain moderate AUC on some days even when F1 at the default threshold is near zero, while the reweighted evaluation shows a meaningful recall difference at a fixed false-positive target. Neither observation licenses a post-hoc threshold chosen on the test set. Instead, threshold selection must be part of the training and validation procedure, and the chosen metric should reflect the analyst workload that alerts will create.")
add_heading(doc, "4.2 Why random can be a strong baseline", 2)
add_text(doc, "Random labeling is sometimes described as a weak baseline because it ignores the current model. At a constrained budget, however, it has two useful properties. It preserves the pool's mixture proportions in expectation and it does not preferentially target a region that the initial model already finds confusing. When the major classes are well represented and the feature space supports a stable linear boundary, those properties can be enough to make the random learner competitive. This is consistent with the UNSW ceiling result, where there is little measurable headroom for any policy.")
add_text(doc, "On CIC, the result is more instructive because there is room to perform worse. All active policies lost to random, including diversity and a genuine vote-disagreement committee. That agreement across policies does not prove that all active learning is harmful, but it weakens explanations that rely on one implementation detail. It also suggests that an evaluation should record not only final F1 but the composition and difficulty of the acquired set. The current artifact preserves the code and seeds needed for such follow-up analysis, while avoiding a causal claim that its measurements do not establish.")
add_heading(doc, "4.3 Dataset integrity and adaptive-policy scope", 2)
add_text(doc, "The CIC results must also be interpreted in light of dataset-construction evidence. Independent audits of CIC-IDS2017 report errors in traffic generation, flow construction, feature extraction, TCP termination, and labeling; corrected or re-extracted versions can materially change benchmark measurements [13,14]. This study uses prepared rows derived from the released CIC flow tables rather than re-extracting every flow from PCAP. Consequently, the day-holdout collapse and cross-dataset failure can reflect an interaction among genuine traffic shift, restricted shared features, model limitations, and flow-exporter artifacts. The experiments do not identify the contribution of each cause.")
add_text(doc, "The active-learning scope is similarly bounded. Recent cross-domain evidence shows that query-strategy rankings vary by domain and seed [12], while NIDS work on model updating considers delayed labels and adaptive strategies in non-stationary streams [15]. The present comparison evaluates four static pool-based policies, not continual, delayed-label, cost-adaptive, or hybrid NIDS systems. A corrected-CIC replication and a streaming adaptive baseline are therefore the most informative extensions; until then, the conclusion applies only to the policies, learners, budgets, and prepared subsets reported here.")
add_text(doc, "Label quality is a plausible contributor to the negative result, not merely a background assumption. Uncertainty sampling deliberately asks for the labels of flows closest to the decision boundary, which are precisely the records most likely to be mislabeled or ambiguous in datasets with known flow-construction and labeling defects [13,14]. If boundary rows carry noisier labels, an acquisition policy that concentrates on them can inherit that noise and actively harm the learner rather than merely fail to help. We do not attribute the CIC deficit to label noise because the released labels do not carry reliability annotations, but the mechanism is consistent with the observed gap and motivates controlled label-noise experiments as future work.")
add_text(doc, "A useful next experiment would compare acquisition policies under controlled changes to label noise, class overlap, and prevalence, then relate each policy's selected examples to those changes. Another would attach realistic annotation costs and adjudication rules to labels. Such work could reveal conditions in which active selection recovers an advantage. Until then, the appropriate default is empirical humility: retain random labeling as a serious baseline and require a repeated, statistically tested gain before treating a query strategy as an efficiency improvement.")

add_heading(doc, "5. Limitations and Ethics", 1)
add_text(doc, "This is an empirical study of prepared public benchmark subsets, not a deployment study or an evaluation of a physical network architecture. The study does not reconstruct each provider's full equipment topology, validate packet-capture placement, or re-extract CIC flows from PCAP with a corrected exporter. The CIC day and UNSW family tests also contain small groups, so several point estimates have wide unreported uncertainty. The cross-dataset test uses only 27 common CICFlowMeter features, and UNSW has no exact feature-name intersection with CIC. Only one 5% prevalence and one 1% false-positive target are evaluated. The committee is a lightweight random-subsample committee, and adaptive or hybrid NIDS-specific acquisition policies are not tested. These limits motivate a narrow conclusion rather than a general claim about all intrusion-detection settings.")
add_text(doc, "All data are public benchmark material and the work is defensive. The artifact contains no live-network scanning or offensive functionality. Any operational adoption would require authorization, privacy review, drift monitoring, and a human process for responding to alerts. The anonymous submission intentionally omits author identity and repository links; the reproducibility package should be shared anonymously only if the target venue permits it.")
add_text(doc, "Several results also depend on choices that were fixed before the final comparisons: the logistic-regression learner, the numeric-feature representation, the batch schedule, and the prepared balanced subsets. A stronger active learner, a representation adapted to encrypted or evolving traffic, or a human-in-the-loop cost model could behave differently. Conversely, the present design intentionally avoids claiming that a small point improvement on a clean benchmark translates into fewer analyst-hours. Future work should evaluate label quality, annotation delay, family novelty, and false-positive cost jointly, ideally on authorized operational data with a documented temporal split.")
add_text(doc, "Finally, the paper uses F1 to make the fixed-budget comparisons legible, even though F1 is not a universal security objective. Different sites may value missed attacks, analyst burden, response latency, or business disruption differently. The operating-point experiment partly addresses this by fixing false-positive rate, but it remains a simplified analysis on a reweighted benchmark subset. Readers should treat the numerical values as reproducible evidence about these protocols, not as a calibrated estimate of a specific organization's risk.")

add_heading(doc, "5.1 Interpretation boundaries", 2)
add_text(doc, "The study also does not claim that the reported random baseline is optimal. Random selection could be improved by stratification using information that would be available before labeling, by a changed seed-set design, or by a workflow that treats uncertainty as a review-prioritization signal rather than a label-acquisition rule. Those variants would be separate policies and should be compared using the same fixed-budget, paired protocol. The negative result is specifically about the four policies implemented here, with the specified flow features, learner, batches, and benchmark pools.")
add_text(doc, "Nor should the results be interpreted as a ranking of the public datasets themselves. The apparent difficulty of CIC relative to UNSW in the active-learning comparison may reflect class construction, feature distributions, collection process, or the interaction of those factors with logistic regression. The strict evaluations help reveal that sensitivity but do not identify a ground-truth source. A causal account would require controlled data generation or richer labels describing ambiguity and annotation reliability; neither is available in the current artifact.")
add_text(doc, "The three benchmarks were chosen because they are flow-based, publicly released with per-row labels and metadata sufficient for day and family splits, and widely used as points of comparison in prior active-learning-for-NIDS work. They are all collected in controlled testbeds and predate recent traffic mixes; the paper does not claim they represent modern production networks. Newer collections (for example CIC-IDS2019 or NF-UNSW-NB15-v2) and real-world telemetry are not used here because they either lack the metadata needed for the strict splits, are less commonly used as active-learning baselines, or were not publicly accessible under compatible terms at artifact time. Extending to them is a straightforward but separate study, and the present claims do not depend on those newer benchmarks being absent.")
add_text(doc, "The active-learning comparisons use balanced 50/50 subsets, which deliberately isolates the acquisition rule from class-imbalance confounding. That choice means the results do not directly test active learning on heavily imbalanced pools, where uncertainty and coverage policies behave differently and where a naive random pool draw would mostly return benign rows. The separate 5% operating-point experiment in Section 3.4 addresses imbalance as a deployment condition, but not active learning under that imbalance; testing acquisition under a native class mix is a distinct future experiment that the current artifact does not claim to cover.")
add_text(doc, "The fixed 100-label budget is likewise a scope decision, not a claim that the comparison holds at every budget. The artifact includes a multi-budget sensitivity harness (20, 50, 100, 240 labels) for this reason, so a reviewer can check whether the random-vs-active ordering persists at intermediate budgets. Within the frozen protocol, 100 labels is the reported budget; the budget sensitivity results are reproduced by the named command in Section 6.1 and are reported here only at the primary budget to keep the tables legible.")
add_text(doc, "These boundaries are deliberate. They turn an otherwise broad claim about active learning into one that can be checked, reproduced, and challenged. A subsequent study can replace the learner, add a cost model, or introduce a new query rule and ask whether it produces a corrected, repeated gain over random sampling under a specified shift. That is a more useful path than interpreting a single benchmark curve as evidence that active learning has solved the labeling problem for intrusion detection.")

add_heading(doc, "5.2 Threats to reproducibility", 2)
add_text(doc, "Empirical reproducibility can fail even when the model code is small. Different data-download mirrors, parser versions, random-number defaults, or silent removal of metadata columns can change a benchmark experiment. The companion artifact therefore pins requirements, retains lightweight metadata subsets, records split definitions and seeds, and provides a checklist of exact reproduction commands. Tests cover the query-policy distinctions and the evaluation guards that are most likely to change the study's interpretation.")
add_text(doc, "Reproducibility does not imply that every rerun will yield identical hardware-level timing or that a different library release will report the same floating-point value. The relevant standard here is stronger than a screenshot but narrower than a production audit: an independent reader should be able to reconstruct the comparisons, observe the direction and statistical status of the active-learning result, and identify where a changed dependency or dataset version alters a conclusion. Any public release should archive the precise dataset preparation instructions allowed by the dataset licenses.")
add_text(doc, "The anonymized paper separates these provenance details from identity. Before double-blind submission, the artifact and supplementary material should be checked for author names, local machine paths, account names, revision history, and repository links. That process is also useful outside peer review: provenance that is explicit and machine-readable makes it easier for later users to distinguish a regenerated result from a result that was copied or manually edited.")

add_heading(doc, "6. Reproducibility", 1)
add_text(doc, "The companion artifact freezes data-preparation commands, package versions, random seeds, split definitions, and evaluation commands. It includes metadata-preserving subsets, a leakage guard for likely label columns, 70 tests, and scripts for active-learning statistics, learner/budget sensitivity, strict generalization, cross-dataset transfer, and operating-point evaluation. The artifact is designed so that each table in this paper can be regenerated from a named command. An anonymized repository must remove author identity, machine paths, and application materials before submission.")
add_text(doc, "The artifact checklist distinguishes source-controlled inputs from regenerated outputs. Metadata CSV subsets, scripts, configuration, and tests are versioned; plots, result JSON, and derived tables are intentionally reproducible outputs rather than opaque committed binaries. This separation makes it possible to audit the experiment after changing a dependency while preserving the exact seeds and held-out definitions used for the paper. It also avoids treating a static figure as the only evidence for a numerical claim.")
add_text(doc, "For an archival or double-blind release, the same checklist should be run in a clean environment and its generated outputs compared with the values reported here. Any discrepancy should be described rather than silently normalized away. The paper's claims are narrow enough that this verification is practical: the key question is whether random labeling remains the strongest CIC policy under the frozen protocol and whether the strict evaluations preserve the stated direction of the distribution-shift warning.")

add_heading(doc, "6.1 Repository, data, and reproduction", 1)
add_text(doc, "Repository URL: withheld from this manuscript for double-blind review and available to reviewers through an anonymous artifact when venue policy permits. Entry points: detector.py, benchmark_compare.py, active_learning_experiment.py, active_learning_stats.py, strict_generalization.py, cross_dataset.py, imbalance_eval.py, plus the prepare_unsw_nb15.py / prepare_cic_ids2017.py / scripts/download_datasets.py data pipeline. Frozen seeds, split definitions, and preprocessing are recorded in reproducibility_config.json and requirements-lock.txt.")
add_text(doc, "Data attribution and license. UNSW-NB15 (N. Moustafa & J. Slay, MilCIS 2015), CIC-IDS2017 (I. Sharafaldin, A. H. Lashkari, & A. Ghorbani, ICISSP 2018), and CSE-CIC-IDS2018 (Canadian Institute for Cybersecurity) are public benchmark datasets used here for research evaluation only; each is subject to its author's license terms. The preprocessed subsets, scripts and code in this artifact are released under the repository's Non-Commercial Personal-Use License.")
add_text(doc, "Reproduction commands.")
add_text(doc, "python -m pip install -r requirements-lock.txt; python scripts/download_datasets.py --all; python benchmark_compare.py --input data/unsw_nb15_subset_with_family.csv --label-column label; python active_learning_stats.py --input data/cic_ids2017_subset_with_day.csv --label-column Label --budget 100 --seeds 8; python active_learning_sensitivity.py --input data/cic_ids2017_subset_with_day.csv --label-column Label --budgets 20 50 100 240 --seeds 8 --output results/cic_active_learning_model_budget_sensitivity.json; python strict_generalization.py --input data/unsw_nb15_subset_with_family.csv --mode family --family-column family --label-column label; python cross_dataset.py --train data/cic_ids2017_subset_with_day.csv --test data/cse_cic_ids2018_subset.csv --label-column Label; python imbalance_eval.py --input data/cic_ids2017_subset_with_day.csv --label-column Label --attack-frac 0.05 --fpr 0.01; python -m pytest -q")

add_heading(doc, "6.2 AI-use disclosure", 1)
add_text(doc, "AI coding assistance was used during implementation and drafting. The author directed the research question, the benchmark evaluation protocol, the strict split and calibration design, the interpretation of the negative result, and reviewed and verified the final code and manuscript claims. AI assistance did not set the research direction or the claims.")

add_heading(doc, "7. Conclusion", 1)
add_text(doc, "Active learning is often treated as an automatic answer to limited labels in security telemetry. Across the prepared benchmarks studied here, it is not. Random labeling matches all alternatives on separable UNSW-NB15 data and significantly outperforms posterior uncertainty, diversity, and QBC on CIC-IDS2017. Stricter day, cross-dataset, and operating-point evaluations further show that impressive balanced random-CV scores can be poor guides to distribution-shift performance. The practical lesson is modest but useful: evaluate the acquisition policy, the split, and the operating point together before claiming that active learning reduces the real labeling burden of intrusion detection.")

add_heading(doc, "References", 1)
for ref in [
    "[1] N. Moustafa and J. Slay. UNSW-NB15: A Comprehensive Data Set for Network Intrusion Detection Systems. MilCIS, 2015.",
    "[2] I. Sharafaldin, A. H. Lashkari, and A. Ghorbani. Toward Generating a New Intrusion Detection Dataset and Intrusion Traffic Characterization. ICISSP, 2018.",
    "[3] Canadian Institute for Cybersecurity. CSE-CIC-IDS2018 Dataset.",
    "[4] F. T. Liu, K. M. Ting, and Z.-H. Zhou. Isolation Forest. ICDM, 2008.",
    "[5] M. Breunig, H.-P. Kriegel, R. Ng, and J. Sander. LOF: Identifying Density-Based Local Outliers. SIGMOD, 2000.",
    "[6] B. Schölkopf et al. Estimating the Support of a High-Dimensional Distribution. Neural Computation, 2001.",
    "[7] D. D. Lewis and W. A. Gale. A Sequential Algorithm for Training Text Classifiers. SIGIR, 1994.",
    "[8] B. Settles. Active Learning Literature Survey. University of Wisconsin-Madison, 2009.",
    "[9] N. Roy and A. McCallum. Toward Optimal Active Learning through Sampling Estimation of Error Reduction. ICML, 2001.",
    "[10] J. Platt. Probabilistic Outputs for Support Vector Machines and Comparisons to Regularized Likelihood Methods. 1999.",
    "[11] A. Niculescu-Mizil and R. Caruana. Predicting Good Probabilities with Supervised Learning. ICML, 2005. doi:10.1145/1102351.1102430.",
    "[12] T. Werner, J. Burchert, M. Stubbemann, and L. Schmidt-Thieme. A Cross-Domain Benchmark for Active Learning. NeurIPS 2024 Datasets and Benchmarks Track. doi:10.52202/079017-2009.",
    "[13] G. Engelen, V. Rimmer, and W. Joosen. Troubleshooting an Intrusion Detection Dataset: the CICIDS2017 Case Study. IEEE Security and Privacy Workshops, 2021. doi:10.1109/SPW53761.2021.00009.",
    "[14] A. Rosay, E. Cheval, F. Carlier, and P. Leroux. Network Intrusion Detection: A Comprehensive Analysis of CIC-IDS2017. ICISSP, 2022. doi:10.5220/0010774000003120.",
    "[15] G. Olímpio, L. Camargos, R. S. Miani, and E. R. Faria. Model Update for Intrusion Detection: Analyzing Delayed Labeling and Active Learning Strategies. Computers & Security 134, 2023. doi:10.1016/j.cose.2023.103451.",
]:
    add_reference(doc, ref)

doc.core_properties.author = "Anonymous"
doc.core_properties.title = "Anonymous Submission"
doc.core_properties.subject = "Network intrusion detection active learning"
doc.core_properties.comments = "Anonymous workshop draft"
doc.save(OUT)
print(OUT)
